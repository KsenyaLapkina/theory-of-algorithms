from clothing.clothing_items import Jacket, Pants
from clothing.suit import ThreePieceSuit
from docx import Document
from openpyxl import Workbook
from database import ClothingDatabase
import os

def show_welcome():
    print("=" * 60)
    print("          КАЛЬКУЛЯТОР ПОШИВА ОДЕЖДЫ")
    print("=" * 60)
    print()

def main():
    # Инициализация базы данных
    db = ClothingDatabase()
    
    show_welcome()
    
    # Ввод данных для пиджака
    print("Параметры пиджака:")
    jacket_size = float(input("Размер пиджака: "))
    has_liner = input("Есть подклад? (да/нет): ").lower() == 'да'
    buttons = int(input("Количество пуговиц: "))
    
    # Ввод данных для брюк
    print("\nПараметры брюк:")
    pants_size = float(input("Размер брюк: "))
    print("Тип брюк: classic (классические), slim (зауженные), wide (широкие), sport (спортивные)")
    pants_type = input("Выберите тип брюк: ").lower()
    has_loops = input("Есть шлёвки для ремня? (да/нет): ").lower() == 'да'
    
    # Выбор формата экспорта
    print("\nВыберите формат сохранения результатов:")
    print("1 - Только Word документ")
    print("2 - Только Excel таблица") 
    print("3 - Оба формата")
    
    export_choice = input("Введите номер выбора (1-3): ")
    export_formats = {
        '1': 'word',
        '2': 'excel', 
        '3': 'both'
    }
    export_format = export_formats.get(export_choice, 'both')
    
    # Создание объектов
    jacket = Jacket(size=jacket_size, has_liner=has_liner, buttons_count=buttons)
    pants = Pants(size=pants_size, pants_type=pants_type, has_belt_loops=has_loops)
    suit = ThreePieceSuit(jacket, pants)
    
    # Вывод результатов
    print("\n" + "="*50)
    print("РЕЗУЛЬТАТЫ РАСЧЕТА:")
    print(f"Пиджак: {jacket}")
    print(f"Стоимость пошива пиджака: {jacket.calculate_sewing_cost():.2f} руб.")
    print(f"Брюки: {pants}")
    print(f"Стоимость пошива брюк: {pants.calculate_sewing_cost():.2f} руб.")
    print(f"Костюм-тройка: {suit.calculate_total_cost():.2f} руб.")
    
    # Сохранение в выбранные форматы
    saved_files = []
    
    if export_format in ['word', 'both']:
        save_to_docx(jacket, pants, suit)
        saved_files.append("'clothing_report.docx'")
    
    if export_format in ['excel', 'both']:
        save_to_excel(jacket, pants, suit)
        saved_files.append("'clothing_report.xlsx'")
    
    # Сохранение в базу данных
    calculation_id = db.save_calculation(jacket, pants, suit, export_format)
    print(f"\nРасчет сохранен в базу данных (ID: {calculation_id})")
    
    print(f"\nОтчеты сохранены в файлы: {', '.join(saved_files)}")
    
    # Показываем историю
    show_history(db)

def show_history(db):
    """Показывает историю расчетов"""
    print("\n" + "="*50)
    print("ПОСЛЕДНИЕ РАСЧЕТЫ:")
    history = db.get_calculation_history()[:3]  # Последние 3 расчета
    
    if not history:
        print("История расчетов пуста")
        return
    
    for calc in history:
        calc_id, timestamp, jacket_size, pants_size, total_cost, export_format = calc
        date = timestamp.split('T')[0]
        print(f"ID: {calc_id} | Дата: {date} | Пиджак: {jacket_size} | Брюки: {pants_size} | Итого: {total_cost:.2f} руб.")

def save_to_docx(jacket, pants, suit):
    """Сохранение отчета в Word"""
    doc = Document()
    doc.add_heading('Отчет о расчете стоимости одежды', 0)
    
    doc.add_heading('Пиджак', level=1)
    doc.add_paragraph(str(jacket))
    doc.add_paragraph(f"Стоимость пошива: {jacket.calculate_sewing_cost():.2f} руб.")
    
    doc.add_heading('Брюки', level=1)
    doc.add_paragraph(str(pants))
    doc.add_paragraph(f"Стоимость пошива: {pants.calculate_sewing_cost():.2f} руб.")
    
    doc.add_heading('Костюм-тройка', level=1)
    doc.add_paragraph(f"Общая стоимость: {suit.calculate_total_cost():.2f} руб.")  # Убрано "со скидкой"
    
    doc.save('clothing_report.docx')

def save_to_excel(jacket, pants, suit):
    """Сохранение отчета в Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Расчет стоимости одежды"
    
    # Заголовки
    headers = ['Изделие', 'Размер', 'Тип', 'Расход ткани (м²)', 'Стоимость ткани (руб.)', 
               'Стоимость работы (руб.)', 'Стоимость фурнитуры (руб.)', 'Общая стоимость (руб.)']
    
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)
    
    # Данные пиджака
    ws.cell(row=2, column=1, value='Пиджак')
    ws.cell(row=2, column=2, value=jacket.size)
    ws.cell(row=2, column=3, value=f"Подклад: {'да' if jacket.has_liner else 'нет'}")
    ws.cell(row=2, column=4, value=jacket.calculate_fabric_consumption())
    ws.cell(row=2, column=5, value=jacket.calculate_fabric_cost())
    ws.cell(row=2, column=6, value=1500.0)  # Понижено с 2000
    ws.cell(row=2, column=7, value=jacket.buttons_count * 50)  # Понижено с 100
    ws.cell(row=2, column=8, value=jacket.calculate_sewing_cost())
    
    # Данные брюк
    ws.cell(row=3, column=1, value='Брюки')
    ws.cell(row=3, column=2, value=pants.size)
    type_names = {"classic": "классические", "slim": "зауженные", "wide": "широкие", "sport": "спортивные"}
    pants_type_name = type_names.get(pants.pants_type, pants.pants_type)
    ws.cell(row=3, column=3, value=pants_type_name)
    ws.cell(row=3, column=4, value=pants.calculate_fabric_consumption())
    ws.cell(row=3, column=5, value=pants.calculate_fabric_cost())
    ws.cell(row=3, column=6, value=1200.0)  # Понижено с 1500
    ws.cell(row=3, column=7, value=200 if pants.has_belt_loops else 0)  # Понижено с 300
    ws.cell(row=3, column=8, value=pants.calculate_sewing_cost())
    
    # Данные костюма
    ws.cell(row=4, column=1, value='Костюм-тройка')
    ws.cell(row=4, column=2, value='Комплект')
    ws.cell(row=4, column=3, value='Полный комплект')  # Убрано "со скидкой"
    ws.cell(row=4, column=8, value=suit.calculate_total_cost())
    
    # Итог
    ws.cell(row=6, column=1, value='ВСЕГО:')
    ws.cell(row=6, column=8, value=suit.calculate_total_cost())
    
    wb.save('clothing_report.xlsx')

if __name__ == "__main__":
    main()